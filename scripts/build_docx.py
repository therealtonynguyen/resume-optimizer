#!/usr/bin/env python3
"""
Generate DOCX from docs/resume.md
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
from docx import Document
from config_loader import get_path, get_output_filename

ROOT = Path(__file__).resolve().parents[1]
SRC = get_path('resume_source')
OUTDIR = get_path('build_dir')
OUTDIR.mkdir(parents=True, exist_ok=True)
OUT = OUTDIR / get_output_filename('baseline_docx')

def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for line in lines:
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.lstrip().startswith("- "):
            doc.add_paragraph(line.strip()[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(str(OUT))
    print(f"Wrote {OUT}")

if __name__ == "__main__":
    main()
